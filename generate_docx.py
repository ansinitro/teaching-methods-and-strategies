from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

data = [
    {
        "title": "The rise of ChatGPT: Exploring its potential in medical education",
        "authors": "Hyunsu Lee (2023)",
        "content": "This article discusses how artificial intelligence, especially ChatGPT, can change medical education. The author explains that AI can act as a virtual teaching assistant to help students learn better. However, teachers must also be careful about ethical problems and update their teaching methods to use this new technology safely.\n\nhttps://doi.org/10.1002/ase.2270"
    },
    {
        "title": "Artificial Intelligence for Assessment and Feedback to Enhance Student Success in Higher Education",
        "authors": "Monika Hooda, et al. (2022)",
        "content": "This study explores how giving students fast and helpful feedback using artificial intelligence improves their learning. The researchers compared different AI methods to see which one works best for grading students. The results help university teachers understand how to use data to create better, less stressful tests for their classes.\n\nhttps://doi.org/10.1155/2022/5215722"
    },
    {
        "title": "ChatGPT Challenges Blended Learning Methodologies in Engineering Education: A Case Study in Mathematics",
        "authors": "Luis Manuel S\u00e1nchez Ruiz, et al. (2023)",
        "content": "This research looks at how engineering students use ChatGPT in their math classes. After surveying 110 students, the authors found that students use the AI tool very often and find it helpful. However, the study warns that teachers need to change their teaching strategies so students still learn important problem-solving skills.\n\nhttps://doi.org/10.3390/app13106039"
    },
    {
        "title": "Artificial Intelligence Technology Assisted Music Teaching Design",
        "authors": "Dan Dai (2021)",
        "content": "This paper explains how artificial intelligence can make teaching music much more effective. By using smart technology and data analysis, teachers can create personalized lessons for each student. The author believes that modernizing music classrooms with AI helps students learn better and makes teaching more interesting.\n\nhttps://doi.org/10.1155/2021/9141339"
    },
    {
        "title": "Challenge, integration, and change: ChatGPT and future anatomical education",
        "authors": "Lige Leng (2024)",
        "content": "This article explores the use of ChatGPT in teaching anatomy to medical students. The author shares their own teaching experience, showing that AI increases student interest and helps them study independently. Despite some challenges, the paper concludes that AI tools are a great addition to modern medical education.\n\nhttps://doi.org/10.1080/10872981.2024.2304973"
    },
    {
        "title": "Artificial intelligence in mathematics education: A systematic literature review",
        "authors": "Riyan Hidayat, et al. (2022)",
        "content": "This review analyzes 20 different studies to understand how AI is used to teach mathematics. The researchers found that AI is mainly used through educational robots, smart software, and virtual teachers. The study shows both the advantages and disadvantages of using these new technologies in math classes.\n\nhttps://doi.org/10.29333/iejme/12132"
    },
    {
        "title": "The application of artificial intelligence assistant to deep learning in teachers' teaching and students' learning processes",
        "authors": "Yi Liu, et al. (2022)",
        "content": "This study investigates how AI assistants help university teachers and students achieve deeper learning. Using questionnaires, the researchers found that AI helps teachers prepare better lessons and makes student learning more personalized. The authors recommend using smart platforms to improve the overall quality of education.\n\nhttps://doi.org/10.3389/fpsyg.2022.929175"
    },
    {
        "title": "Evaluation of Online Teaching Quality of Basic Education Based on Artificial Intelligence",
        "authors": "Moyan Li, Yawen Su (2020)",
        "content": "This paper focuses on the challenge of keeping online classes high quality. The researchers developed a new AI model to evaluate and grade how well online teaching works. Based on their results, they suggest several new strategies that schools can use to improve their digital education programs.\n\nhttps://doi.org/10.3991/ijet.v15i16.15937"
    },
    {
        "title": "Generative AI and Higher Education: Trends, Challenges, and Future Directions from a Systematic Literature Review",
        "authors": "Jo\u00e3o Batista, et al. (2024)",
        "content": "This review examines 37 recent studies about how generative AI is changing universities. The findings show that while students like using AI, there are still major challenges with cheating and academic rules. The authors suggest that universities need to create better strategies for teaching and testing students in the future.\n\nhttps://doi.org/10.3390/info15110676"
    },
    {
        "title": "Using Artificial Intelligence for Developing English Language Teaching/Learning: An Analytical Study from University Students\u2019 Perspective",
        "authors": "Turki Rabah Al Mukhallafi (2020)",
        "content": "This study explores what university students think about using artificial intelligence to learn English. After surveying 44 students, the researcher found that while AI is very effective, it is currently not being used enough in classrooms. The paper suggests a detailed plan for teachers to include more AI activities in English courses.\n\nhttps://doi.org/10.5539/ijel.v10n6p40"
    }
]

doc = Document()

p_student = doc.add_paragraph()
p_student.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_student.add_run('Astana IT University\n').bold = True
p_student.add_run('Program: Applied Artificial Intelligence\n')
p_student.add_run('Course: Teaching Methods and Strategies\n')
p_student.add_run('Instructor: Peter Shon\n\n')
p_student.add_run('FullName: Angsar Shaumen\n')
p_student.add_run('Group: AAI-2501M\n')
p_student.add_run('ID: 255782\n')
p_student.add_run('E-Mail: 255782@astanait.edu.kz\n\n')

heading = doc.add_heading('Annotated bibliographic list of literature on the topic "Innovative Teaching Methods using Artificial Intelligence in Higher Education"', level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in heading.runs:
    run.font.size = Pt(14)
    run.font.bold = True

p1 = doc.add_paragraph()
p1.add_run('Keywords: ').bold = True
p1.add_run('«Artificial Intelligence in Education», «Innovative Teaching Methods», «Teaching Strategies», «Generative AI», «Digital Learning», «Higher Education Technology»')

p2 = doc.add_paragraph()
p2.add_run('Databases: ').bold = True
p2.add_run('OpenAlex, Scopus, Web of Science, Crossref')

p3 = doc.add_paragraph()
p3.add_run('Range: ').bold = True
p3.add_run('2020 to 2025')

p4 = doc.add_paragraph()
p4.add_run('Quantity: ').bold = True
p4.add_run('at least 10 articles (All Open Access - Free to read!)')

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.autofit = False
table.columns[0].width = Inches(0.5)
table.columns[1].width = Inches(2.5)
table.columns[2].width = Inches(1.5)
table.columns[3].width = Inches(3.5)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Article'
hdr_cells[1].text = 'title'
hdr_cells[2].text = 'Author (s)'
hdr_cells[3].text = 'Brief abstract, link to the source'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, item in enumerate(data):
    row_cells = table.add_row().cells
    row_cells[0].text = f"{i+1}."
    row_cells[1].text = item['title']
    row_cells[2].text = item['authors']
    
    content_parts = item['content'].split('\n\n')
    p = row_cells[3].paragraphs[0]
    p.add_run(content_parts[0] + '\n\n')
    link_run = p.add_run(content_parts[1])
    link_run.font.color.rgb = RGBColor(0, 0, 255)
    link_run.font.underline = True

doc.save('c:/Users/ansinitro/Desktop/teaching/Annotated_Bibliography_AI_Teaching_Final.docx')
print("Document generated successfully at: c:/Users/ansinitro/Desktop/teaching/Annotated_Bibliography_AI_Teaching_Final.docx")
